"""
json_to_word.py
───────────────
Reconstructs a Word (.docx) from a JSON file produced by word_to_html.py.
Faithfully restores: fonts, sizes, bold/italic/underline, colors, alignment,
indentation, spacing, tab stops, tables (borders + shading + column widths),
and text boxes (positioned, bordered).

Output goes to:  to_word/<basename>.docx
"""

import json
import os
import copy
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree

# ── Twips helper ──────────────────────────────────────────────────────────────
def _tw(v, default=0):
    try: return Twips(int(v))
    except: return Twips(default)

def _twips_to_emu(v):
    try: return int(int(v) * 914400 / 1440)
    except: return 0

# ── Color helper ──────────────────────────────────────────────────────────────
def _rgb(hex_str):
    if not hex_str: return None
    h = hex_str.lstrip("#")
    if len(h) != 6: return None
    try:
        return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))
    except: return None

# ── Alignment map ─────────────────────────────────────────────────────────────
_ALIGN = {
    "left":    WD_ALIGN_PARAGRAPH.LEFT,
    "center":  WD_ALIGN_PARAGRAPH.CENTER,
    "right":   WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "both":    WD_ALIGN_PARAGRAPH.JUSTIFY,
    "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
}

# ── Document view setting (Print Layout) ─────────────────────────────────────
def _set_print_layout(doc):
    """Force Print Layout view when document opens."""
    settings = doc.settings.element
    view_el = OxmlElement("w:view")
    view_el.set(qn("w:val"), "print")
    for old in settings.findall(qn("w:view")):
        settings.remove(old)
    settings.insert(0, view_el)


def _patch_doc_defaults(doc, defaults: dict):
    """
    Overwrite the w:pPrDefault spacing inside w:docDefaults in styles.xml.
    """
    styles_el = doc.styles.element

    doc_defs = styles_el.find(qn("w:docDefaults"))
    if doc_defs is None:
        doc_defs = OxmlElement("w:docDefaults")
        styles_el.insert(0, doc_defs)

    ppr_default = doc_defs.find(qn("w:pPrDefault"))
    if ppr_default is None:
        ppr_default = OxmlElement("w:pPrDefault")
        doc_defs.append(ppr_default)

    ppr = ppr_default.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        ppr_default.append(ppr)

    for old in ppr.findall(qn("w:spacing")):
        ppr.remove(old)

    spacing = OxmlElement("w:spacing")

    after  = defaults.get("default_space_after")
    before = defaults.get("default_space_before")
    ls     = defaults.get("default_line_spacing")
    lsrule = defaults.get("default_line_spacing_rule", "auto")

    spacing.set(qn("w:after"),  str(int(after  if after  is not None else 0)))
    spacing.set(qn("w:before"), str(int(before if before is not None else 0)))

    if ls is not None:
        spacing.set(qn("w:line"), str(int(ls)))
        rule_map = {"exact": "exact", "atLeast": "atLeast"}
        spacing.set(qn("w:lineRule"), rule_map.get(lsrule, "auto"))

    ppr.append(spacing)

# ── Page setup ────────────────────────────────────────────────────────────────
def _apply_page(section, page_meta):
    section.page_width    = _tw(page_meta.get("width",  12240))
    section.page_height   = _tw(page_meta.get("height", 15840))
    section.top_margin    = _tw(page_meta.get("margin_top",    1440))
    section.bottom_margin = _tw(page_meta.get("margin_bottom", 1440))
    section.left_margin   = _tw(page_meta.get("margin_left",   1800))
    section.right_margin  = _tw(page_meta.get("margin_right",  1800))

# ── Run formatting ─────────────────────────────────────────────────────────────
def _apply_run(run, fmt, defaults=None):
    defaults = defaults or {}

    font_name    = fmt.get("font")         or defaults.get("font")    or "Calibri"
    font_cs      = fmt.get("font_cs")     or defaults.get("font_cs") or font_name
    font_eastAsia= fmt.get("font_eastAsia")
    font_size    = fmt.get("size")        or defaults.get("size")    or 11
    font_size_cs = fmt.get("size_cs")     or font_size
    bold      = fmt.get("bold")
    italic    = fmt.get("italic")
    underline = fmt.get("underline")
    strike    = fmt.get("strike")
    color     = fmt.get("color")
    rtl       = fmt.get("rtl", False)

    run.font.name = font_name
    run.font.size = Pt(float(font_size))

    rPr = run._r.get_or_add_rPr()

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"),  font_name)
    rFonts.set(qn("w:hAnsi"),  font_name)
    rFonts.set(qn("w:cs"),     font_cs)
    if font_eastAsia:
        rFonts.set(qn("w:eastAsia"), font_eastAsia)

    for old_el in rPr.findall(qn("w:szCs")):
        rPr.remove(old_el)
    szCs_el = OxmlElement("w:szCs")
    szCs_el.set(qn("w:val"), str(int(float(font_size_cs) * 2)))
    rPr.append(szCs_el)

    if bold is not None:
        run.font.bold = bold
        if bold:
            rPr = run._r.get_or_add_rPr()
            for old_el in rPr.findall(qn("w:bCs")):
                rPr.remove(old_el)
            rPr.append(OxmlElement("w:bCs"))
    if italic is not None:
        run.font.italic = italic
        if italic:
            rPr = run._r.get_or_add_rPr()
            for old_el in rPr.findall(qn("w:iCs")):
                rPr.remove(old_el)
            rPr.append(OxmlElement("w:iCs"))
    if underline is not None: run.font.underline = underline
    if strike    is not None: run.font.strike    = strike
    if color:
        rgb = _rgb(color)
        if rgb: run.font.color.rgb = rgb

    char_spacing = fmt.get("char_spacing")
    if char_spacing is not None:
        rPr = run._r.get_or_add_rPr()
        for old_el in rPr.findall(qn("w:spacing")):
            rPr.remove(old_el)
        sp_el = OxmlElement("w:spacing")
        sp_el.set(qn("w:val"), str(int(char_spacing)))
        rPr.append(sp_el)

    rPr = run._r.get_or_add_rPr()
    for old_el in rPr.findall(qn("w:rtl")):
        rPr.remove(old_el)
    if rtl:
        rPr.append(OxmlElement("w:rtl"))

    for old_el in rPr.findall(qn("w:cs")):
        rPr.remove(old_el)
    if rtl and font_name:
        rPr.append(OxmlElement("w:cs"))

# ── Paragraph formatting ───────────────────────────────────────────────────────
def _apply_paragraph_fmt(para, pdata, defaults=None):
    pf  = para.paragraph_format
    cfg = pdata

    raw_align = cfg.get("align")
    if raw_align:
        align = _ALIGN.get(raw_align)
        if align is not None:
            para.alignment = align

    if "space_before" in cfg:
        pf.space_before = _tw(cfg["space_before"])
    if "space_after" in cfg:
        pf.space_after  = _tw(cfg["space_after"])
    if "indent_left" in cfg:
        pf.left_indent  = _tw(cfg["indent_left"])
    if "indent_right" in cfg:
        pf.right_indent = _tw(cfg["indent_right"])
    if "indent_first" in cfg:
        pf.first_line_indent = _tw(cfg["indent_first"])
    if "indent_hanging" in cfg:
        pf.first_line_indent = _tw(-cfg["indent_hanging"])

    if "line_spacing" in cfg:
        pf.line_spacing = _tw(cfg["line_spacing"])
        rule = cfg.get("line_spacing_rule","auto")
        if rule == "exact":
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        elif rule == "atLeast":
            pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        else:
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    for ts in cfg.get("tab_stops",[]):
        _add_tab_stop(para, ts["pos"], ts.get("type","left"), ts.get("leader","none"))

    if cfg.get("keep_lines"): pf.keep_together = True
    if cfg.get("keep_next"):  pf.keep_with_next = True
    if cfg.get("page_break_before"):
        pf.page_break_before = True

    if cfg.get("contextual_spacing"):
        pPr = para._p.get_or_add_pPr()
        for old in pPr.findall(qn("w:contextualSpacing")):
            pPr.remove(old)
        pPr.append(OxmlElement("w:contextualSpacing"))

    if cfg.get("before_autospacing") or cfg.get("after_autospacing"):
        pPr = para._p.get_or_add_pPr()
        sp_el = pPr.find(qn("w:spacing"))
        if sp_el is None:
            sp_el = OxmlElement("w:spacing")
            pPr.append(sp_el)
        if cfg.get("before_autospacing"):
            sp_el.set(qn("w:beforeAutospacing"), "1")
        if cfg.get("after_autospacing"):
            sp_el.set(qn("w:afterAutospacing"), "1")

    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:bidi")):
        pPr.remove(old)
    bidi_el = OxmlElement("w:bidi")
    if not cfg.get("bidi"):
        bidi_el.set(qn("w:val"), "0")   # explicit LTR — overrides any inherited RTL
    pPr.append(bidi_el)

    pmark = cfg.get("pmark_fmt", {})
    if pmark:
        pPr = para._p.get_or_add_pPr()
        for old_rpr in pPr.findall(qn("w:rPr")):
            pPr.remove(old_rpr)
        pm_rPr = OxmlElement("w:rPr")

        pm_font      = pmark.get("font")
        pm_font_cs   = pmark.get("font_cs") or pm_font
        pm_eastAsia  = pmark.get("font_eastAsia")
        if pm_font or pm_font_cs:
            pm_rFonts = OxmlElement("w:rFonts")
            if pm_font:     pm_rFonts.set(qn("w:ascii"),   pm_font)
            if pm_font:     pm_rFonts.set(qn("w:hAnsi"),   pm_font)
            if pm_font_cs:  pm_rFonts.set(qn("w:cs"),      pm_font_cs)
            if pm_eastAsia: pm_rFonts.set(qn("w:eastAsia"),pm_eastAsia)
            pm_rPr.append(pm_rFonts)

        if pmark.get("bold"):
            pm_rPr.append(OxmlElement("w:b"))
            pm_rPr.append(OxmlElement("w:bCs"))
        if pmark.get("italic"):
            pm_rPr.append(OxmlElement("w:i"))
            pm_rPr.append(OxmlElement("w:iCs"))

        pm_size    = pmark.get("size")
        pm_size_cs = pmark.get("size_cs") or pm_size
        if pm_size:
            sz_el = OxmlElement("w:sz")
            sz_el.set(qn("w:val"), str(int(float(pm_size) * 2)))
            pm_rPr.append(sz_el)
        if pm_size_cs:
            szCs_el = OxmlElement("w:szCs")
            szCs_el.set(qn("w:val"), str(int(float(pm_size_cs) * 2)))
            pm_rPr.append(szCs_el)

        if pmark.get("char_spacing") is not None:
            sp_el = OxmlElement("w:spacing")
            sp_el.set(qn("w:val"), str(int(pmark["char_spacing"])))
            pm_rPr.append(sp_el)

        if pmark.get("underline"):
            u_el = OxmlElement("w:u")
            u_el.set(qn("w:val"), "single")
            pm_rPr.append(u_el)

        if pmark.get("color"):
            rgb = pmark["color"].lstrip("#")
            if len(rgb) == 6:
                c_el = OxmlElement("w:color")
                c_el.set(qn("w:val"), rgb.upper())
                pm_rPr.append(c_el)

        if pmark.get("rtl"):
            pm_rPr.append(OxmlElement("w:rtl"))

        if len(pm_rPr):
            pPr.append(pm_rPr)

def _add_tab_stop(para, pos_twips, kind="left", leader="none"):
    pPr  = para._p.get_or_add_pPr()
    tabs_el = pPr.find(qn("w:tabs"))
    if tabs_el is None:
        tabs_el = OxmlElement("w:tabs")
        pPr.append(tabs_el)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), kind)
    tab.set(qn("w:pos"), str(int(pos_twips)))
    if leader and leader != "none":
        tab.set(qn("w:leader"), leader)
    tabs_el.append(tab)

# ── Style ID → python-docx name map ──────────────────────────────────────────
# IMPORTANT: styleId in OOXML uses CamelCase with no spaces ("TableGrid"),
# but python-docx looks up styles by their display name ("Table Grid").
# Missing entries here cause style_spacing overrides to fail silently,
# which lets docDefaults (e.g. space_after=200 twips = 10pt) bleed into
# table cell paragraphs.
_STYLE_ID_MAP = {
    "Normal":      "Normal",
    "Heading1":    "Heading 1",
    "Heading2":    "Heading 2",
    "Heading3":    "Heading 3",
    "Heading4":    "Heading 4",
    "NoSpacing":   "No Spacing",
    "BodyText":    "Body Text",
    "Caption":     "Caption",
    "Header":      "Header",
    "Footer":      "Footer",
    # ── Table styles (styleId has no space; python-docx name has space) ──
    "TableGrid":    "Table Grid",
    "TableNormal":  "Table Normal",
    "LightShading": "Light Shading",
    "LightList":    "Light List",
    "LightGrid":    "Light Grid",
}

def _resolve_style(doc, style_id):
    if not style_id:
        return None
    mapped = _STYLE_ID_MAP.get(style_id, style_id)
    try:
        return doc.styles[mapped]
    except KeyError:
        try:
            return doc.styles[style_id]
        except KeyError:
            return None

# ── Add paragraph to document/cell ────────────────────────────────────────────
def _add_paragraph(container, pdata, defaults=None):
    runs = pdata.get("runs", [])

    if not runs and pdata.get("text","").strip():
        runs = [{"text": pdata["text"]}]

    if hasattr(container, "add_paragraph"):
        style_id = pdata.get("style")
        style_name = _STYLE_ID_MAP.get(style_id, style_id) if style_id else None
        try:
            para = container.add_paragraph(style=style_name)
        except (KeyError, Exception):
            para = container.add_paragraph()
    else:
        para = container
        para.clear()

    _apply_paragraph_fmt(para, pdata, defaults)

    for run_data in runs:
        text = run_data.get("text","")
        if "\n" in text:
            parts = text.split("\n")
            for idx, part in enumerate(parts):
                if idx > 0:
                    run = para.add_run()
                    run.add_break()
                if part:
                    run = para.add_run(part)
                    _apply_run(run, run_data, defaults)
        elif "\t" in text:
            parts = text.split("\t")
            last_run = None
            for idx, part in enumerate(parts):
                if part:
                    last_run = para.add_run(part)
                    _apply_run(last_run, run_data, defaults)
                if idx < len(parts) - 1:
                    if last_run is None:
                        last_run = para.add_run()
                        _apply_run(last_run, run_data, defaults)
                    t_el = OxmlElement("w:tab")
                    last_run._r.append(t_el)
                    last_run = None
        else:
            run = para.add_run(text)
            _apply_run(run, run_data, defaults)

    return para

# ── Border XML helper ─────────────────────────────────────────────────────────
def _border_el(tag, bdata):
    if not bdata: return None
    el = OxmlElement(tag)
    el.set(qn("w:val"),   bdata.get("style","single"))
    el.set(qn("w:sz"),    str(bdata.get("size",4)))
    color = bdata.get("color","000000")
    if color: el.set(qn("w:color"), color.lstrip("#"))
    el.set(qn("w:space"), str(bdata.get("space",0)))
    return el

def _apply_cell_border(cell, borders_dict):
    if not borders_dict: return
    tc  = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    bEl = OxmlElement("w:tcBorders")
    for side in ("top","bottom","left","right","insideH","insideV"):
        if side in borders_dict:
            sub = _border_el(f"w:{side}", borders_dict[side])
            if sub is not None:
                bEl.append(sub)
    tcPr.append(bEl)

def _apply_cell_bg(cell, hex_color):
    if not hex_color: return
    tc   = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#").upper())
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    tcPr.append(shd)

def _apply_cell_width(cell, width_twips):
    if not width_twips: return
    tc   = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(int(width_twips)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.insert(0, tcW)

def _apply_cell_vmerge(cell, value):
    tc   = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    for old in tcPr.findall(qn("w:vMerge")):
        tcPr.remove(old)
    vm = OxmlElement("w:vMerge")
    if value == "restart":
        vm.set(qn("w:val"), "restart")
    tcPr.append(vm)

# ── Table ─────────────────────────────────────────────────────────────────────
def _add_table(doc, tdata, defaults=None):
    rows = tdata.get("rows", [])
    if not rows: return

    col_widths = tdata.get("col_widths", [])
    tbl_width  = tdata.get("width", 0)

    # col_widths is the most reliable source for true grid column count.
    # Using max(len(cells)) is WRONG for tables with colspan — a row with
    # 4 cell entries where one has colspan=2 actually fills 5 grid columns.
    if col_widths:
        n_cols = len(col_widths)
    else:
        # Count effective columns by summing colspans per row
        n_cols = max(
            (sum(c.get("colspan", 1) for c in r.get("cells", []))
             for r in rows),
            default=1
        )
        if tbl_width and n_cols:
            col_widths = [tbl_width // n_cols] * n_cols

    tbl = doc.add_table(rows=0, cols=n_cols)
    tbl.style = "Table Grid"

    bidi_visual = tdata.get("bidi_visual", False)

    tbl_el = tbl._tbl
    tblPr  = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)

    ins_pos = 0

    if bidi_visual:
        for old in tblPr.findall(qn("w:bidiVisual")):
            tblPr.remove(old)
        tblPr.insert(ins_pos, OxmlElement("w:bidiVisual"))
        ins_pos += 1

    if tbl_width:
        for old in tblPr.findall(qn("w:tblW")):
            tblPr.remove(old)
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"),    str(int(tbl_width)))
        tblW.set(qn("w:type"), "dxa")
        tblPr.insert(ins_pos, tblW)
        ins_pos += 1

    tbl_align = tdata.get("align")
    if tbl_align and not bidi_visual:
        jc_el = OxmlElement("w:jc")
        jc_el.set(qn("w:val"), tbl_align)
        tblPr.insert(ins_pos, jc_el)

    if col_widths:
        for old in tbl_el.findall(qn("w:tblGrid")):
            tbl_el.remove(old)
        grid = OxmlElement("w:tblGrid")
        for w in col_widths:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(int(w)))
            grid.append(gc)
        tbl_el.append(grid)

    for row_data in rows:
        cells_data = row_data.get("cells", [])
        tr_el = tbl.add_row()
        doc_cells = tr_el.cells

        if "height" in row_data:
            trPr = tr_el._tr.get_or_add_trPr()
            trH  = OxmlElement("w:trHeight")
            trH.set(qn("w:val"),    str(row_data["height"]))
            trH.set(qn("w:hRule"), row_data.get("height_rule","auto"))
            trPr.append(trH)

        col_cursor = 0
        for ci, cell_data in enumerate(cells_data):
            if col_cursor >= n_cols: break
            cell = doc_cells[col_cursor]

            cw = cell_data.get("width") or (col_widths[col_cursor] if col_cursor < len(col_widths) else 0)
            if cw: _apply_cell_width(cell, cw)

            if cell_data.get("bg"):
                _apply_cell_bg(cell, cell_data["bg"])

            if cell_data.get("borders"):
                _apply_cell_border(cell, cell_data["borders"])

            if cell_data.get("vmerge"):
                _apply_cell_vmerge(cell, cell_data["vmerge"])

            cs = cell_data.get("colspan", 1)
            if cs > 1:
                end = min(col_cursor + cs - 1, n_cols - 1)
                cell.merge(doc_cells[end])

            paras = cell_data.get("paragraphs", [])
            if not paras:
                p = cell.paragraphs[0]
                p.clear()
                # No data at all → zero out spacing so docDefaults don't bleed in
                p.paragraph_format.space_after  = Pt(0)
                p.paragraph_format.space_before = Pt(0)
            else:
                first = True
                for pdata in paras:
                    target = cell.paragraphs[0] if first else cell.add_paragraph()
                    p = _add_paragraph(target, pdata, defaults)
                    first = False
                    # Table cells: if no explicit spacing in the JSON, zero it out.
                    # Without this, paragraphs inherit docDefaults (e.g. 200 twips = 10pt)
                    # which is correct for body text but wrong for table cells.
                    if "space_after" not in pdata:
                        p.paragraph_format.space_after  = Pt(0)
                    if "space_before" not in pdata:
                        p.paragraph_format.space_before = Pt(0)

            col_cursor += cs

        # If this row's cells don't fill all columns, merge the last real
        # cell rightward to cover remaining columns — prevents phantom empty columns.
        if col_cursor < n_cols and col_cursor > 0:
            last_real_cell = doc_cells[col_cursor - 1]
            tail_cell      = doc_cells[n_cols - 1]
            if last_real_cell._tc is not tail_cell._tc:
                last_real_cell.merge(tail_cell)

    tbl_borders = tdata.get("borders")
    if tbl_borders:
        tbl_el = tbl._tbl
        tblPr  = tbl_el.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl_el.insert(0, tblPr)
        for old in tblPr.findall(qn("w:tblBorders")):
            tblPr.remove(old)
        tbBord = OxmlElement("w:tblBorders")
        for side in ("top","bottom","left","right","insideH","insideV"):
            if side in tbl_borders:
                sub = _border_el(f"w:{side}", tbl_borders[side])
                if sub is not None:
                    tbBord.append(sub)
        tblPr.append(tbBord)

    doc.add_paragraph()

# ── Text box (floating, anchored) ─────────────────────────────────────────────
def _add_textbox(doc, tbdata, defaults=None, host_para=None):
    pos_x  = int(tbdata.get("pos_x",  0))
    pos_y  = int(tbdata.get("pos_y",  0))
    width  = int(tbdata.get("width",  1440))
    height = int(tbdata.get("height", 1440))

    pos_x_emu  = _twips_to_emu(pos_x)
    pos_y_emu  = _twips_to_emu(pos_y)
    width_emu  = _twips_to_emu(width)
    height_emu = _twips_to_emu(height)

    pos_h_anchor = tbdata.get("pos_h_anchor", "page")
    pos_v_anchor = tbdata.get("pos_v_anchor", "page")

    has_border = tbdata.get("has_border", True)
    border_xml = (
        '<a:ln xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" w="12700">'
        '<a:solidFill><a:srgbClr val="000000"/></a:solidFill></a:ln>'
        if has_border else
        '<a:ln xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" w="0">'
        '<a:noFill/></a:ln>'
    )

    paras_xml = ""
    for pdata in tbdata.get("paragraphs", []):
        paras_xml += _paragraph_to_xml(pdata, defaults)

    run_xml = f"""<w:r
     xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
     xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
     xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
     xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
     xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
    <w:rPr/>
    <w:drawing>
      <wp:anchor distT="0" distB="0" distL="114300" distR="114300"
                 simplePos="0" relativeHeight="251659264" behindDoc="0"
                 locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="{pos_h_anchor}">
          <wp:posOffset>{pos_x_emu}</wp:posOffset>
        </wp:positionH>
        <wp:positionV relativeFrom="{pos_v_anchor}">
          <wp:posOffset>{pos_y_emu}</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="{width_emu}" cy="{height_emu}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="1" name="TextBox 1"/>
        <wp:cNvGraphicFramePr/>
        <a:graphic>
          <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
            <wps:wsp>
              <wps:cNvSpPr txBx="1">
                <a:spLocks noChangeArrowheads="1"/>
              </wps:cNvSpPr>
              <wps:spPr>
                <a:xfrm>
                  <a:off x="{pos_x_emu}" y="{pos_y_emu}"/>
                  <a:ext cx="{width_emu}" cy="{height_emu}"/>
                </a:xfrm>
                <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                <a:noFill/>
                {border_xml}
              </wps:spPr>
              <wps:txbx>
                <w:txbxContent>
                  {paras_xml}
                </w:txbxContent>
              </wps:txbx>
              <wps:bodyPr rot="0" spcFirstLastPara="0" vertOverflow="overflow"
                          horzOverflow="overflow" vert="horz" wrap="square"
                          lIns="91440" tIns="45720" rIns="91440" bIns="45720"
                          anchor="t" anchorCtr="0"/>
            </wps:wsp>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>
  </w:r>"""

    run_el = etree.fromstring(run_xml)

    if host_para is not None:
        host_para._p.append(run_el)
    else:
        host = doc.add_paragraph()
        host._p.append(run_el)

# ── Paragraph → XML string (for text box content) ────────────────────────────
def _paragraph_to_xml(pdata, defaults=None):
    ns  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    pEl = etree.Element(f"{{{ns}}}p")

    pPr = etree.SubElement(pEl, f"{{{ns}}}pPr")

    align = pdata.get("align")
    if align:
        jc = etree.SubElement(pPr, f"{{{ns}}}jc")
        jc.set(f"{{{ns}}}val", align)

    if pdata.get("bidi"):
        etree.SubElement(pPr, f"{{{ns}}}bidi")

    ind_l = pdata.get("indent_left")
    ind_r = pdata.get("indent_right")
    ind_h = pdata.get("indent_hanging")
    ind_f = pdata.get("indent_first")
    if any(v is not None for v in (ind_l, ind_r, ind_h, ind_f)):
        ind_el = etree.SubElement(pPr, f"{{{ns}}}ind")
        if ind_l is not None: ind_el.set(f"{{{ns}}}left",      str(int(ind_l)))
        if ind_r is not None: ind_el.set(f"{{{ns}}}right",     str(int(ind_r)))
        if ind_h is not None: ind_el.set(f"{{{ns}}}hanging",   str(int(ind_h)))
        if ind_f is not None: ind_el.set(f"{{{ns}}}firstLine", str(int(ind_f)))

    tab_stops = pdata.get("tab_stops", [])
    if tab_stops:
        tabs_el = etree.SubElement(pPr, f"{{{ns}}}tabs")
        for ts in tab_stops:
            tab_el = etree.SubElement(tabs_el, f"{{{ns}}}tab")
            tab_el.set(f"{{{ns}}}val", ts.get("type", "left"))
            tab_el.set(f"{{{ns}}}pos", str(int(ts.get("pos", 0))))

    for run_data in pdata.get("runs", []):
        rEl  = etree.SubElement(pEl, f"{{{ns}}}r")
        rPr  = etree.SubElement(rEl, f"{{{ns}}}rPr")

        font_name = run_data.get("font") or (defaults or {}).get("font", "Calibri")
        font_cs   = run_data.get("font_cs") or font_name
        rFonts = etree.SubElement(rPr, f"{{{ns}}}rFonts")
        rFonts.set(f"{{{ns}}}ascii",  font_name)
        rFonts.set(f"{{{ns}}}hAnsi",  font_name)
        rFonts.set(f"{{{ns}}}cs",     font_cs)

        sz_val    = run_data.get("size")    or (defaults or {}).get("size", 11)
        sz_cs_val = run_data.get("size_cs") or sz_val
        sz = etree.SubElement(rPr, f"{{{ns}}}sz")
        sz.set(f"{{{ns}}}val", str(int(float(sz_val) * 2)))
        szCs = etree.SubElement(rPr, f"{{{ns}}}szCs")
        szCs.set(f"{{{ns}}}val", str(int(float(sz_cs_val) * 2)))

        if run_data.get("bold"):
            etree.SubElement(rPr, f"{{{ns}}}b")
            etree.SubElement(rPr, f"{{{ns}}}bCs")
        if run_data.get("italic"):
            etree.SubElement(rPr, f"{{{ns}}}i")
            etree.SubElement(rPr, f"{{{ns}}}iCs")
        if run_data.get("underline"):
            u = etree.SubElement(rPr, f"{{{ns}}}u")
            u.set(f"{{{ns}}}val", "single")
        if run_data.get("color"):
            rgb_str = run_data["color"].lstrip("#")
            if len(rgb_str) == 6:
                c_el = etree.SubElement(rPr, f"{{{ns}}}color")
                c_el.set(f"{{{ns}}}val", rgb_str.upper())
        if run_data.get("rtl"):
            etree.SubElement(rPr, f"{{{ns}}}rtl")
            etree.SubElement(rPr, f"{{{ns}}}cs")

        text = run_data.get("text", "")
        if "\t" in text:
            parts = text.split("\t")
            for idx, part in enumerate(parts):
                if part:
                    tEl = etree.SubElement(rEl, f"{{{ns}}}t")
                    tEl.text = part
                    if part[0] == " " or part[-1] == " ":
                        tEl.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                if idx < len(parts) - 1:
                    etree.SubElement(rEl, f"{{{ns}}}tab")
        else:
            tEl = etree.SubElement(rEl, f"{{{ns}}}t")
            tEl.text = text
            if text and (text[0] == " " or text[-1] == " "):
                tEl.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    return etree.tostring(pEl, encoding="unicode")

# ── Image placeholder ─────────────────────────────────────────────────────────
def _add_image_placeholder(doc, img_data, defaults=None):
    target = img_data.get("target","")
    w = img_data.get("width",0); h = img_data.get("height",0)
    note = f"[IMAGE: {os.path.basename(target) or 'unknown'} {w}×{h} twips]"
    p = doc.add_paragraph(note)
    for run in p.runs:
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x99,0x99,0x99)

# ── Main reconstruct function ──────────────────────────────────────────────────
def reconstruct_docx(json_path):
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    meta     = data.get("meta", {})
    page     = meta.get("page", {})
    defaults = meta.get("defaults", {"font":"Calibri","size":11})

    doc = Document()
    _set_print_layout(doc)

    normal_style = doc.styles["Normal"]
    normal_pf    = normal_style.paragraph_format
    normal_pf.line_spacing        = None
    normal_pf.line_spacing_rule   = WD_LINE_SPACING.SINGLE
    normal_pf.space_before        = Pt(0)
    normal_pf.space_after         = Pt(0)

    if doc.paragraphs:
        pPr = doc.paragraphs[0]._p.get_or_add_pPr()
        for old in pPr.findall(qn("w:spacing")):
            pPr.remove(old)
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:before"), "0")
        sp.set(qn("w:after"),  "0")
        pPr.append(sp)

    if defaults.get("default_space_before") is not None:
        normal_pf.space_before = _tw(defaults["default_space_before"])
    if defaults.get("default_space_after") is not None:
        normal_pf.space_after  = _tw(defaults["default_space_after"])
    if defaults.get("default_line_spacing") is not None:
        normal_pf.line_spacing = _tw(defaults["default_line_spacing"])
        rule = defaults.get("default_line_spacing_rule","auto")
        if rule == "exact":
            normal_pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        elif rule == "atLeast":
            normal_pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        else:
            normal_pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    _patch_doc_defaults(doc, defaults)

    # ── Apply named style spacing overrides ───────────────────────────────────
    # KEY FIX: Use the full _STYLE_ID_MAP (including TableGrid → "Table Grid")
    # so that style_spacing["TableGrid"] = {space_after: 0} actually finds the
    # "Table Grid" style and writes the override. Without this, the docDefaults
    # value of 200 twips (10pt) bleeds into every table cell paragraph.
    style_spacing = defaults.get("style_spacing", {})
    if style_spacing:
        _style_by_id   = {s.style_id: s for s in doc.styles}
        _style_by_name = {s.name:     s for s in doc.styles}
        for style_id, sdata in style_spacing.items():
            st = _style_by_id.get(style_id) or _style_by_name.get(style_id)
            if st is None:
                # Use the full map (includes TableGrid → "Table Grid")
                mapped = _STYLE_ID_MAP.get(style_id)
                if mapped:
                    st = _style_by_name.get(mapped)
            if st is None:
                continue
            spf = st.paragraph_format
            if "space_before" in sdata: spf.space_before = _tw(sdata["space_before"])
            if "space_after"  in sdata: spf.space_after  = _tw(sdata["space_after"])
            if "line_spacing" in sdata:
                spf.line_spacing = _tw(sdata["line_spacing"])
                rule = sdata.get("line_spacing_rule","auto")
                if rule == "exact":
                    spf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                elif rule == "atLeast":
                    spf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                else:
                    spf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    _apply_page(doc.sections[0], page)

    last_para = None

    for el in data.get("elements", []):
        etype = el.get("type","")

        if etype == "paragraph":
            last_para = _add_paragraph(doc, el, defaults)

        elif etype == "empty_paragraph":
            style_id = el.get("style")
            style_name = _STYLE_ID_MAP.get(style_id, style_id) if style_id else None
            try:
                p = doc.add_paragraph(style=style_name)
            except (KeyError, Exception):
                p = doc.add_paragraph()
            _apply_paragraph_fmt(p, el, defaults)
            last_para = p

        elif etype == "table":
            _add_table(doc, el, defaults)
            last_para = None

        elif etype == "textbox":
            _add_textbox(doc, el, defaults, host_para=last_para)

        elif etype == "image":
            _add_image_placeholder(doc, el, defaults)
            last_para = None

    return doc

# ── Convenience entry point ───────────────────────────────────────────────────
def reconstruct_to_file(json_path, out_dir="to_word"):
    os.makedirs(out_dir, exist_ok=True)
    basename = os.path.splitext(os.path.basename(json_path))[0]
    out_path = os.path.join(out_dir, basename + ".docx")
    doc = reconstruct_docx(json_path)
    doc.save(out_path)
    return out_path

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python json_to_word.py <file.json> [out_dir]")
        sys.exit(1)
    out_dir = sys.argv[2] if len(sys.argv) > 2 else "to_word"
    result  = reconstruct_to_file(sys.argv[1], out_dir)
    print(f"Reconstructed → {result}")